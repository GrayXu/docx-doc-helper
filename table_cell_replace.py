'''
2021/7/1
author: grayxu

批量替换表格内的内容，同时不改变表格文本的格式。
'''

import docx
from tqdm import tqdm

in_path = '/home/gray/desktop/7.1.docx'
out_path = '/home/gray/desktop/out.docx'
# in_path = '/home/gray/desktop/test2.docx'
# out_path = '/home/gray/desktop/out2.docx'

doc = docx.Document(in_path)

# generate codes
pair = {'必选的':['是','强制的','强制','强制性的','强制性','必需的','必选'],
       '有条件的':['条件选择','条件','条件的'],
       '可选的':['否','可选','任选的','任选'],
       '要求':['需求']}

replace = {}
for key in pair.keys():
    for v in pair[key]:
        replace[v] = [key]

print(replace)

# content replace
for table_index in tqdm(range(len(list(doc.tables)))):
    table = doc.tables[table_index]
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:  # use cell paragraph to keep text style
                data = []
                for run in paragraph.runs:
                    data.append(run.text.strip())
                cell_text = ''.join(data)

                if cell_text in replace.keys():
                    isFirst = True
                    for run in paragraph.runs:
                        if isFirst:
                            run.text = replace[cell_text]
                            isFirst = False
                        else:
                            run.text = ''

doc.save(out_path)