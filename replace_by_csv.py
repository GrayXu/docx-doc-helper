'''
2021/7/5
author: grayxu

根据CSV表格内提交的信息来批量替换原文同section内的关键词。
CSV格式：
1（序号）,23.1.2.1 简介（出处，提供section位置）,（原文关键词）,（新关键词）,（提交机构）,（采纳|拒绝）,

note: 表格等序号不统一问题，暂时仅支持正文
'''

import csv
import docx
import utils
from tqdm import tqdm


# config
in_path = '/home/gray/desktop/7.1.docx'
out_path = '/home/gray/desktop/out.docx'
csv_path = '/home/gray/desktop/dell.csv'

doc = docx.Document(in_path)

replace = {}
with open(csv_path) as file:
    data = csv.reader(file,delimiter=',')
    count = 0
    for row in data: 
        count += 1
        if count <= 3:
            continue
        
        sec_num = utils.get_sec_num(row[1])       
        
        if sec_num is not None:
            # tmp solution: self-increment for a new added-in section
            tmp_s = [int(x) for x in sec_num.split('.')]
            tmp_s[0] += 1
            
            sec_num = ".".join([str(x) for x in tmp_s])
            if sec_num not in replace.keys():
                replace[sec_num] = []
            replace[sec_num].append((row[2],row[3]))
            
#         replace[] = (row[2],row[3])

doc = docx.Document(in_path)

found = {}

# search in paragraphs
now_sec = ''
for para in tqdm(list(doc.paragraphs)):
    # update sec_num
    sec_num = utils.get_sec_num(para.text)
    if sec_num is not None:
        now_sec = sec_num
    
    if now_sec == '':
        continue
    
    if now_sec not in replace.keys():
        continue
    
    for a,b in replace[now_sec]:
#         for run in para.runs:
#             if a != '' and a in run.text:
#                 if a not in found.keys():
#                     found[a] = 0
#                 found[a] += 1

        if a != '' and a in para.text:
            '''
            replace(a,b)
            '''
            if a not in found.keys():
                found[a] = 0
            found[a] += 1


print(found)