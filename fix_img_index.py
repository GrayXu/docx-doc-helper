'''
2021/7/13
author: grayxu
设置所有图片的style，自动编号

note: 图片旧标题均为Caption Style，开头为图xx，后跟空格或\u3000或直接连接文字。
'''

import docx
import re

doc = docx.Document('/home/gray/desktop/7.12.docx')

for para in doc.paragraphs:
#     if para.text.startswith('图'):
#         print(para.text)
    result = re.match('图\d+ ', para.text)
    if result is not None:
        print(para.text)
    else:
        result = re.match('图\d+\u3000', para.text)
        if result is not None:
            print(para.text)
            
            
# 批量处理不是图片title的style
# 1. 通过正则删除编号信息
# 2. 设置style
pattern = re.compile(r'图\d+\s*')
for para in doc.paragraphs:
    if get_style_name(para.style) == 'Caption' and para.text.startswith('图'):
        para.style = '正文图标题'  # 黑体 5号
        pattern.findall(target)
        para_text = re.sub(pattern, '',para.text)
        
        # clear
        for r in para.runs:
            r.text = ''
        para.runs[-1].text = para_text

doc.save('/home/gray/desktop/output.docx')